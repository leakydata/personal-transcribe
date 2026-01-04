"""
Word document (DOCX) exporter for PersonalTranscribe.
Generates editable Word documents with formatted transcripts.
"""

import os
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from src.models.transcript import Transcript, format_timestamp, format_timestamp_range
from src.export.base_exporter import BaseExporter


class DOCXExporter(BaseExporter):
    """Exports transcripts to Word document format."""
    
    @property
    def format_name(self) -> str:
        return "Word Document"
    
    @property
    def file_extension(self) -> str:
        return ".docx"
    
    def export(
        self,
        transcript: Transcript,
        output_path: str,
        audio_file: Optional[str] = None,
        include_timestamps: bool = True,
        include_line_numbers: bool = True,
        include_gaps: bool = True,
        include_header: bool = True,
        font_size: int = 11,
        certification_text: str = "",
        metadata = None,
        **options
    ) -> None:
        """Export transcript to DOCX.
        
        Args:
            transcript: Transcript to export
            output_path: Path for output DOCX
            audio_file: Source audio filename (for header)
            include_timestamps: Include timestamp for each segment
            include_line_numbers: Include line numbers
            include_gaps: Show gap indicators between segments
            include_header: Include document header
            font_size: Base font size in points
            certification_text: Optional legal certification text
            metadata: Optional RecordingMetadata for header
        """
        output_path = self.validate_output_path(output_path)
        self.ensure_output_directory(output_path)
        
        # Create document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(font_size)
        
        # Recording metadata section (if provided)
        if metadata and not metadata.is_empty():
            self._add_metadata_section(doc, metadata)
        
        # Document title
        if include_header:
            title = doc.add_heading('Transcript', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Basic metadata (if no full metadata provided)
            if not metadata or metadata.is_empty():
                meta = doc.add_paragraph()
                meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if audio_file:
                    meta.add_run(f'Source: {os.path.basename(audio_file)}\n')
                
                meta.add_run(f'Date: {datetime.now().strftime("%Y-%m-%d %H:%M")}\n')
                meta.add_run(f'Duration: {format_timestamp(transcript.audio_duration)}\n')
                meta.add_run(f'Segments: {transcript.segment_count} | Words: {transcript.word_count}')
            
            doc.add_paragraph()  # Spacer
        
        # Create table for transcript
        num_cols = 1  # Always have text column
        if include_line_numbers:
            num_cols += 1
        if include_timestamps:
            num_cols += 1
        
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        hdr_cells = table.rows[0].cells
        col_idx = 0
        
        if include_line_numbers:
            hdr_cells[col_idx].text = '#'
            hdr_cells[col_idx].paragraphs[0].runs[0].bold = True
            col_idx += 1
        
        if include_timestamps:
            hdr_cells[col_idx].text = 'Time'
            hdr_cells[col_idx].paragraphs[0].runs[0].bold = True
            col_idx += 1
        
        hdr_cells[col_idx].text = 'Text'
        hdr_cells[col_idx].paragraphs[0].runs[0].bold = True
        
        # Gap threshold for display
        gap_threshold = 2.0  # Show gaps >= 2 seconds
        
        # Add transcript segments
        for i, segment in enumerate(transcript.segments):
            # Calculate gap before this segment
            gap_before = 0.0
            if i == 0:
                gap_before = segment.start_time
            else:
                prev_segment = transcript.segments[i - 1]
                gap_before = segment.start_time - prev_segment.end_time
            
            # Add gap indicator BEFORE segment if significant
            if include_gaps and gap_before >= gap_threshold:
                gap_row = table.add_row().cells
                
                # Merge all cells for gap indicator
                gap_cell = gap_row[0]
                for j in range(1, num_cols):
                    gap_cell = gap_cell.merge(gap_row[j])
                
                gap_para = gap_cell.paragraphs[0]
                gap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if gap_before >= 60:
                    mins = int(gap_before // 60)
                    secs = int(gap_before % 60)
                    gap_text = f'[ PAUSE: {mins}m {secs}s - Other party speaking / Silence ]'
                else:
                    gap_text = f'[ PAUSE: {gap_before:.0f} seconds - Other party speaking / Silence ]'
                
                gap_run = gap_para.add_run(gap_text)
                gap_run.font.bold = True
                gap_run.font.italic = True
                gap_run.font.color.rgb = RGBColor(70, 130, 180)  # Steel blue
            
            # Add segment row
            row_cells = table.add_row().cells
            col_idx = 0
            
            if include_line_numbers:
                row_cells[col_idx].text = str(i + 1)
                col_idx += 1
            
            if include_timestamps:
                time_str = format_timestamp_range(segment.start_time, segment.end_time)
                time_para = row_cells[col_idx].paragraphs[0]
                time_run = time_para.add_run(time_str)
                time_run.font.name = 'Consolas'
                time_run.font.size = Pt(font_size - 1)
                time_run.font.color.rgb = RGBColor(100, 100, 100)
                col_idx += 1
            
            # Text cell
            text_para = row_cells[col_idx].paragraphs[0]
            text_para.add_run(segment.display_text)
            
            # Mark bookmarked segments
            if segment.is_bookmarked:
                for run in text_para.runs:
                    run.font.highlight_color = 6  # Yellow highlight
        
        # Set column widths
        if include_line_numbers:
            table.columns[0].width = Inches(0.4)
        if include_timestamps:
            col_idx = 1 if include_line_numbers else 0
            table.columns[col_idx].width = Inches(1.5)
        
        # Certification text
        if certification_text:
            doc.add_paragraph()  # Spacer
            doc.add_heading('Certification', level=1)
            doc.add_paragraph(certification_text)
            
            # Signature lines
            doc.add_paragraph()
            sig_para = doc.add_paragraph()
            sig_para.add_run('_' * 40 + '\n')
            sig_para.add_run('Signature')
            
            doc.add_paragraph()
            date_para = doc.add_paragraph()
            date_para.add_run('_' * 40 + '\n')
            date_para.add_run('Date')
        
        # Save document
        doc.save(output_path)
    
    def _add_metadata_section(self, doc, metadata) -> None:
        """Add recording metadata section to document.
        
        Args:
            doc: Document object
            metadata: RecordingMetadata object
        """
        # Title
        heading = doc.add_heading('Recording Information', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Case information
        if metadata.case_name or metadata.case_number:
            p = doc.add_paragraph()
            p.add_run('Case/Matter: ').bold = True
            if metadata.case_name:
                p.add_run(f'{metadata.case_name}\n')
            if metadata.case_number:
                p.add_run(f'Case No: {metadata.case_number}\n')
            if metadata.client_name:
                p.add_run(f'Client: {metadata.client_name}')
        
        # Recording details
        p = doc.add_paragraph()
        p.add_run('Recording Details:\n').bold = True
        
        if metadata.recording_date:
            date_str = metadata.recording_date
            if metadata.recording_time:
                date_str += f' at {metadata.recording_time}'
            p.add_run(f'Date/Time: {date_str}\n')
        
        if metadata.recording_location:
            p.add_run(f'Location: {metadata.recording_location}\n')
        
        if metadata.recording_source:
            p.add_run(f'Source: {metadata.recording_source}\n')
        
        if metadata.original_filename:
            p.add_run(f'File: {metadata.original_filename}\n')
        
        if metadata.audio_duration:
            p.add_run(f'Duration: {metadata.format_duration()}')
        
        # Participants
        if metadata.participants:
            p = doc.add_paragraph()
            p.add_run('Participants:\n').bold = True
            for i, participant in enumerate(metadata.participants, 1):
                p.add_run(f'{i}. {participant}\n')
        
        # Transcription info
        if metadata.transcriptionist or metadata.transcription_date:
            p = doc.add_paragraph()
            p.add_run('Transcription:\n').bold = True
            if metadata.transcriptionist:
                p.add_run(f'Transcribed by: {metadata.transcriptionist}\n')
            if metadata.transcription_date:
                p.add_run(f'Date: {metadata.transcription_date}')
        
        # Notes
        if metadata.notes:
            p = doc.add_paragraph()
            p.add_run('Notes:\n').bold = True
            run = p.add_run(metadata.notes)
            run.italic = True
        
        # Divider
        doc.add_paragraph('_' * 60)
        doc.add_paragraph()